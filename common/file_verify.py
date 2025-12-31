# -*- coding: utf-8 -*-
"""
@Project : InterfaceTest
@File    : file_verify.py
@Author  : Chlon
@Date    : 2025/12/29 14:10
@Desc    : 文件验证
"""
import io
import requests
from docx import Document
import pdfplumber

class FileVerify:

    @staticmethod
    def verify_file_content(file_url, expected_keywords):
        """
        验证文件内容是否包含指定的关键词
        :param file_url: 文件下载地址
        :param expected_keywords: 期望存在的关键词列表，如 ["上海交通大学", "2025年"]
        """
        response = requests.get(file_url)
        assert response.status_code == 200

        # 将二进制内容转为字节流，方便库读取
        file_stream = io.BytesIO(response.content)
        full_text = ""

        # === 分支 A: 验证 Word (.docx) ===
        if file_url.endswith(".docx"):
            try:
                doc = Document(file_stream)
                # 提取段落文字
                paras = [p.text for p in doc.paragraphs]
                # 提取表格文字 (报告中很多数据在表格里)
                tables = []
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            tables.append(cell.text)

                full_text = "\n".join(paras + tables)
            except Exception as e:
                raise AssertionError(f"Word文件解析失败，可能文件已损坏: {e}")

        # === 分支 B: 验证 PDF ===
        elif file_url.endswith(".pdf"):
            try:
                with pdfplumber.open(file_stream) as pdf:
                    # 遍历每一页提取文字
                    pages_text = [page.extract_text() for page in pdf.pages]
                    full_text = "\n".join(filter(None, pages_text))
            except Exception as e:
                raise AssertionError(f"PDF文件解析失败: {e}")

        else:
            print(f"⚠️ 未知文件格式，跳过内容验证: {file_url}")
            return

        # === 核心断言 ===
        print(f"正在验证文件内容，长度: {len(full_text)}字符...")

        for keyword in expected_keywords:
            if keyword in full_text:
                print(f"✅ 关键词验证通过: {keyword}")
            else:
                # 打印部分文本方便调试
                print(f"❌ 关键词缺失: {keyword}")
                print(f"文件开头文本: {full_text[:200]}")
                raise AssertionError(f"报告文件中未找到关键词: {keyword}")

    